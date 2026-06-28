r"""Tests for the "forge from an existing sheet" ingestion path.

Two halves, both API-key-free:

1. **Extraction** (`extract_sheet_text`): .docx, .pdf, .txt/.md, the "no text"
   (scanned-image) path, and the unsupported-type error each behave. Fixtures are
   built in-process — a real .docx via python-docx, a minimal text-layer PDF authored
   here so pypdf has something to extract.
2. **Assembly**: feed an extracted sheet's text through `autofill` with a FAKE model
   (mirroring test_autofill_pc) and assert the same {character, warnings} contract the
   front-end already consumes — kind=="character", pc{} populated, docx_text reached the
   prompt.

Run:  .venv_forge\Scripts\python tests\test_sheet_extract.py
"""
from __future__ import annotations

import io
import json
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT))

from forge.agents import autofill                                      # noqa: E402
from forge.agents.sheet_extract import (                              # noqa: E402
    extract_sheet_text, UnsupportedSheetType,
)


# -- fixtures -----------------------------------------------------------------
def _make_docx(paragraphs, table_rows=None) -> bytes:
    import docx

    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                table.cell(r, c).text = val
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _make_pdf(text: str) -> bytes:
    """A minimal single-page PDF with a real text layer, xref offsets computed."""
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length %d >>\nstream\nBT /F1 24 Tf 72 700 Td (%s) Tj ET\nendstream"
        % (len(text) + 26, text.encode("latin-1")),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n%s\nendobj\n" % (i, body)
    xref_pos = len(out)
    out += b"xref\n0 %d\n" % (len(objs) + 1)
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<< /Size %d /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF" % (
        len(objs) + 1, xref_pos)
    return bytes(out)


def _make_blank_pdf() -> bytes:
    """A PDF with a page but no text layer — stands in for a scanned/image-only sheet."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# -- the fake PC model (mirrors test_autofill_pc) -----------------------------
FAKE_PC = {
    "ruleset": "dnd5e-2014", "kind": "character",
    "name": "Thorin Oakenshield", "size": "Medium", "type": "humanoid (dwarf)",
    "alignment": "Lawful Good", "flavour": "A dwarf fighter reclaiming a lost hold.",
    "ac": "18 (chain mail, shield)", "hp": "30 (4d10 + 8)", "speed": "25 ft.",
    "resist": "", "condImm": "", "languages": "Common, Dwarvish",
    "challenge": "— (level 4)",
    "traits": [{"name": "Second Wind", "text": "Regain 1d10+4 HP as a bonus action.", "source": "class:Fighter"}],
    "actions": [{"name": "Battleaxe", "text": "Melee Weapon Attack: +6 to hit. Hit: 1d8 + 4 slashing.", "source": "class:Fighter"}],
    "reactions": [],
    "dump": "",
    "art": {"appearance": "Broad dwarf, braided beard", "outfit": "Chain mail",
            "pose": "Axe raised", "environment": "Mountain hall", "personality": "Proud",
            "style": "Painterly high fantasy"},
    "pc": {
        "species": "dwarf", "subspecies": "hill-dwarf", "class": "fighter",
        "subclass": "champion", "level": 4, "background": "",
        "abilityMethod": "standard_array",
        "baseAbilities": {"STR": 15, "DEX": 10, "CON": 14, "INT": 8, "WIS": 13, "CHA": 12},
        "abilityAllocation2024": "",
        "skillChoices": ["athletics", "intimidation"],
        "hitDice": {"die": "d10", "total": 4, "remaining": 4},
        "deathSaves": {"successes": 0, "failures": 0},
        "feats": [], "equipment": ["Battleaxe", "Chain mail", "Shield"],
        "currency": {"cp": 0, "sp": 0, "ep": 0, "gp": 10, "pp": 0},
        "personality": {"traits": ["Stoic"], "ideals": ["Honor"], "bonds": ["My lost hold"], "flaws": ["Stubborn"]},
    },
    "spellcasting": {"ability": "", "cantrips": [], "prepared": []},
}

SHEET_TEXT_MARKER = "Thorin Oakenshield — Level 4 Dwarf Fighter"


def fake_model(system: str, user: str) -> dict:
    assert "KIND: character" in user
    assert "UPLOADED DOCUMENT" in user        # the extracted sheet text was folded in
    assert SHEET_TEXT_MARKER in user          # ...and it's the text we extracted
    return json.loads(json.dumps(FAKE_PC))    # deep copy


# -- tests --------------------------------------------------------------------
def main() -> int:
    failures: list[str] = []

    # 1. docx: paragraphs + table cells
    docx_bytes = _make_docx(
        ["Character: Thorin Oakenshield", "Class: Fighter 4"],
        table_rows=[["STR", "15"], ["CON", "14"]],
    )
    docx_text = extract_sheet_text("sheet.docx", docx_bytes)
    print(f"docx text:\n  {docx_text!r}")
    if "Thorin Oakenshield" not in docx_text:
        failures.append("docx: paragraph text missing")
    if "STR" not in docx_text or "15" not in docx_text:
        failures.append("docx: table-cell text missing")

    # 2. pdf with a text layer
    pdf_text = extract_sheet_text("sheet.pdf", _make_pdf("Thorin the Dwarf Fighter"))
    print(f"pdf text: {pdf_text!r}")
    if "Thorin" not in pdf_text:
        failures.append(f"pdf: text not extracted, got {pdf_text!r}")

    # 3. txt / md pass-through
    txt_text = extract_sheet_text("notes.txt", "A halfling rogue.".encode("utf-8"))
    if txt_text != "A halfling rogue.":
        failures.append(f"txt: pass-through wrong, got {txt_text!r}")
    md_text = extract_sheet_text("notes.md", b"# Pip\n\nfearless")
    if "Pip" not in md_text:
        failures.append(f"md: pass-through wrong, got {md_text!r}")

    # 4. no-text path: a scanned/image-only PDF extracts to empty
    blank = extract_sheet_text("scanned.pdf", _make_blank_pdf())
    print(f"blank pdf -> {blank!r}")
    if blank != "":
        failures.append(f"blank pdf should extract to '', got {blank!r}")

    # 5. unsupported type raises
    try:
        extract_sheet_text("portrait.png", b"\x89PNG...")
        failures.append("unsupported .png should have raised UnsupportedSheetType")
    except UnsupportedSheetType:
        pass

    # 6. assembly: extracted text -> autofill (fake model) -> {character, warnings}
    sheet_bytes = _make_docx([SHEET_TEXT_MARKER, "Champion subclass; chain mail + shield"])
    text = extract_sheet_text("thorin.docx", sheet_bytes)
    res = autofill("", kind="character", ruleset="dnd5e-2014", docx_text=text, model=fake_model)
    char = res["character"]
    pc = char.get("pc", {})
    print(f"\nassembled: kind={char.get('kind')} class={pc.get('class')} "
          f"challenge={char.get('challenge')!r} warnings={len(res['warnings'])}")
    if char.get("kind") != "character":
        failures.append("assembly: kind should be 'character'")
    if not pc:
        failures.append("assembly: pc{} should be populated")
    if pc.get("class") != "fighter":
        failures.append(f"assembly: pc.class should be fighter, got {pc.get('class')}")
    if "character" not in res or "warnings" not in res:
        failures.append("assembly: result must have {character, warnings} keys")
    if not isinstance(res.get("warnings"), list):
        failures.append("assembly: warnings must be a list")

    print()
    if failures:
        print("FAIL:")
        for f in failures:
            print(f"  - {f}")
        return 1
    print("OK: sheet extraction (docx/pdf/txt/md/no-text/unsupported) + assembly contract all pass.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
