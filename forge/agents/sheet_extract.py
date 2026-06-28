"""Text extraction for the "forge from an existing sheet" path.

`extract_sheet_text(filename, data)` turns an uploaded file's bytes into plain text
the auto-fill agent can read. It dispatches by extension only — the byte payload is
never trusted to "sniff" a format:

- ``.docx`` -> python-docx (paragraphs + every table cell, in reading order)
- ``.pdf``  -> pypdf (concatenate each page's extracted text)
- ``.txt`` / ``.md`` -> decode as UTF-8 (pass-through)

This is a best-effort *draft* feeder, not a parser: the AI reads the text, the engine
validates the numbers, the user fixes the rest in the Studio. Two things are explicitly
out of scope for v1:

- **Unsupported types** raise ``UnsupportedSheetType`` (the endpoint maps it to a 400).
- **Scanned / image-only PDFs** (no embedded text layer) are NOT OCR'd. They extract to
  an empty string; the caller detects "no text" and surfaces a clean warning rather than
  silently feeding nothing to the agent.
"""
from __future__ import annotations

import io
from pathlib import Path

SUPPORTED_EXTENSIONS = (".docx", ".pdf", ".txt", ".md")


class UnsupportedSheetType(ValueError):
    """Raised when an uploaded file's extension isn't one we can read."""


def extract_sheet_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded sheet's bytes, dispatched by extension.

    Returns the extracted text (possibly empty — e.g. a scanned PDF with no text layer;
    the caller is expected to treat empty as "no readable text"). Raises
    ``UnsupportedSheetType`` for any extension we don't handle.
    """
    ext = Path(filename or "").suffix.lower()
    if ext == ".docx":
        return _extract_docx(data)
    if ext == ".pdf":
        return _extract_pdf(data)
    if ext in (".txt", ".md"):
        return _extract_text(data)
    raise UnsupportedSheetType(
        f"Unsupported file type {ext or '(none)'!r}. "
        f"Upload one of: {', '.join(SUPPORTED_EXTENSIONS)}."
    )


def _extract_docx(data: bytes) -> str:
    """Join every paragraph and every table cell, in document order."""
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    lines: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                lines.append(line)
    return "\n".join(line for line in lines if line.strip()).strip()


def _extract_pdf(data: bytes) -> str:
    """Concatenate each page's extracted text. Empty for image-only (scanned) PDFs."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(t.strip() for t in pages if t.strip()).strip()


def _extract_text(data: bytes) -> str:
    """Decode plain text / markdown as UTF-8, tolerating a stray byte or two."""
    return data.decode("utf-8", errors="replace").strip()
